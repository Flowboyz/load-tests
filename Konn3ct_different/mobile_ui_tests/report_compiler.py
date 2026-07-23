import os
import re
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def strip_ansi_codes(text):
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    return ansi_escape.sub('', text)

def analyze_logs(log_lines):
    analysis = {
        "summary": "The test run completed, but some assertions or steps did not succeed.",
        "category": "General Execution Failure",
        "details": "Please review the raw console logs for more details.",
        "links": [],
        "device": "Unknown Device"
    }
    
    cleaned = [strip_ansi_codes(line).strip() for line in log_lines]
    
    # Look for Maestro cloud links
    link_pattern = re.compile(r'https://app\.maestro\.dev/\S+')
    for line in cleaned:
        found_links = link_pattern.findall(line)
        for l in found_links:
            # Strip trailing braces/quotes if any
            l_clean = l.rstrip(')').rstrip(']').rstrip('"').rstrip("'")
            if l_clean not in analysis["links"]:
                analysis["links"].append(l_clean)
                
    # Look for device specs
    device_flag = False
    device_lines = []
    for line in cleaned:
        if "Maestro cloud device specs:" in line:
            device_flag = True
            continue
        if device_flag:
            if line.startswith("*"):
                device_lines.append(line.replace("*", "").strip())
            elif line.startswith("?"):
                device_flag = False
                
    if device_lines:
        analysis["device"] = ", ".join(device_lines)
        
    full_log_text = "\n".join(cleaned)
    
    if "Cannot enter raw mode on a non-interactive terminal" in full_log_text:
        analysis["category"] = "Interactive Terminal Error"
        analysis["summary"] = "The Maestro CLI crashed because it tried to prompt for project selection in a headless environment."
        analysis["details"] = "This occurs when your Maestro Cloud API Key is associated with multiple projects and the --project-id flag is not provided. To resolve this, ensure you configure the Project ID in the dashboard settings."
    elif "No flows in workspace match app ID" in full_log_text:
        mismatch_match = re.search(r"No flows in workspace match app ID '(.*?)'\. Found app IDs: (.*)", full_log_text)
        analysis["category"] = "App ID Mismatch"
        analysis["summary"] = "Maestro Cloud rejected the test because the appId defined in the test script does not match the package name of the uploaded APK."
        if mismatch_match:
            expected = mismatch_match.group(1)
            found = mismatch_match.group(2)
            analysis["details"] = (
                f"The uploaded APK was packaged under '{expected}', but the test flow was configured for '{found}'.\n"
                f"Recommendation: Update the App ID in your configuration or flow file to '{expected}' to match the APK's true package name."
            )
        else:
            analysis["details"] = "The appId defined in the YAML flow file does not match the uploaded APK's package name."
    elif "Assertion is false" in full_log_text or "Assertion failed" in full_log_text:
        analysis["category"] = "Functional Assertion Failure"
        analysis["summary"] = "A user interface assertion failed during execution (an expected UI element was not visible)."
        assertion_match = re.search(r"Assertion is false:\s*(.*)", full_log_text)
        if not assertion_match:
            assertion_match = re.search(r"Assertion failed:\s*(.*)", full_log_text)
        if assertion_match:
            reason = assertion_match.group(1)
            analysis["details"] = (
                f"Failed assertion: '{reason}'.\n"
                f"This typically means the application did not transition to the expected screen in time, "
                f"or the element ID/text used in the test step does not exist on the current screen.\n"
                f"Recommendation: Verify if the login/onboarding flow of the app matches the selectors in your script."
            )
        else:
            analysis["details"] = "A test assertion failed. An element was not found or was not visible on the screen."
    elif "Maestro execution completed successfully" in full_log_text or "1/1 Flow Passed" in full_log_text:
        analysis["category"] = "Successful Run"
        analysis["summary"] = "The test execution completed successfully. All steps passed!"
        analysis["details"] = "No errors were detected in the log stream."
        
    return analysis

def generate_mobile_reports(flow_path, device_id, log_lines, duration_sec):
    """
    Generates Markdown and Word reports based on Maestro test execution output.
    """
    # 1. Parse flow and device info
    flow_name = os.path.basename(flow_path)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    timestamp_readable = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Clean all log lines of ANSI escape codes first to prevent XML/docx crashes
    cleaned_log_lines = [strip_ansi_codes(line) for line in log_lines]
    
    # 2. Extract steps and statuses from log lines
    steps = []
    verdict = "PASS"
    error_details = None
    
    for line in cleaned_log_lines:
        clean_line = line.strip()
        if not clean_line:
            continue
            
        # Match steps like "🚀 Step 1: launchApp -> [SUCCESS]" or similar
        step_match = re.search(r'Step (\d+):\s*(.*?)\s*->\s*\[(.*?)\]', clean_line)
        if step_match:
            idx = int(step_match.group(1))
            action = step_match.group(2)
            status = step_match.group(3)
            steps.append({
                "index": idx,
                "action": action,
                "status": status
            })
            if status != "SUCCESS":
                verdict = "FAIL"
            continue
            
        # Match real Maestro output: e.g. "  ✓  launchApp (3.2s)" or "  ✗  tapOn Login (5.0s)"
        real_step_match = re.search(r'^([✓✗])\s*(.*?)(?:\s+\((\d+(?:\.\d+)?s)\))?$', clean_line)
        if real_step_match:
            symbol = real_step_match.group(1)
            action = real_step_match.group(2).strip()
            duration = real_step_match.group(3) or "N/A"
            status = "SUCCESS" if symbol == "✓" else "FAILED"
            
            steps.append({
                "index": len(steps) + 1,
                "action": f"{action} ({duration})" if duration != "N/A" else action,
                "status": status
            })
            if status != "SUCCESS":
                verdict = "FAIL"
            continue
        
        # Match failures in Maestro output
        if "❌" in clean_line or "failed" in clean_line.lower() or "error" in clean_line.lower():
            if not error_details:
                error_details = clean_line
            verdict = "FAIL"

    # Post-process: Check for success completion marker or failed marker to override verdict
    for line in cleaned_log_lines:
        clean_line = line.strip()
        if "Maestro execution completed successfully" in clean_line or "1/1 Flow Passed" in clean_line:
            verdict = "PASS"
            break
        elif "Maestro failed" in clean_line or "Maestro crashed" in clean_line or "1/1 Flow Failed" in clean_line:
            verdict = "FAIL"
            break
            
    # Force verdict to FAIL if any step failed
    if any(s["status"] == "FAILED" for s in steps):
        verdict = "FAIL"
                
    # If no steps were parsed (e.g. CLI not found or crash before steps)
    if not steps:
        verdict = "FAIL"
        steps.append({
            "index": 1,
            "action": "Initialize Test Suite",
            "status": "FAILED"
        })
        if not error_details:
            error_details = "Maestro crashed or CLI was not found."
 
    # 3. Create reports directory
    reports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "mobile_reports")
    os.makedirs(reports_dir, exist_ok=True)
    
    md_filename = f"mobile_report_{timestamp}.md"
    docx_filename = f"mobile_report_{timestamp}.docx"
    
    md_filepath = os.path.join(reports_dir, md_filename)
    docx_filepath = os.path.join(reports_dir, docx_filename)
    
    # --- Generate Markdown Report ---
    write_markdown_report(md_filepath, flow_name, device_id, verdict, steps, error_details, duration_sec, timestamp_readable, cleaned_log_lines)
    
    # --- Generate DOCX Report ---
    try:
        write_docx_report(docx_filepath, flow_name, device_id, verdict, steps, error_details, duration_sec, timestamp_readable, cleaned_log_lines)
    except Exception as e:
        print(f"Error compiling DOCX report: {e}")
        
    return {
        "md_name": md_filename,
        "docx_name": docx_filename,
        "md_path": md_filepath,
        "docx_path": docx_filepath
    }

def write_markdown_report(filepath, flow_name, device_id, verdict, steps, error_details, duration_sec, timestamp_str, log_lines):
    verdict_color = "🟢 **PASS**" if verdict == "PASS" else "🔴 **FAIL**"
    analysis_data = analyze_logs(log_lines)
    
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"# Konn3ct Mobile UI Test Report\n\n")
        f.write(f"## 1. Executive Summary\n")
        f.write(f"| Metric | Value |\n")
        f.write(f"| :--- | :--- |\n")
        f.write(f"| **Test Flow** | {flow_name} |\n")
        f.write(f"| **Target Device** | {device_id} |\n")
        f.write(f"| **Execution Date** | {timestamp_str} |\n")
        f.write(f"| **Duration** | {duration_sec:.2f} seconds |\n")
        f.write(f"| **Overall Verdict** | {verdict_color} |\n\n")
        
        if error_details:
            f.write(f"> [!WARNING]\n")
            f.write(f"> **Error Details**: {error_details}\n\n")
            
        f.write(f"## 2. Step Execution Breakdown\n")
        f.write(f"| Step | Action / Command | Status |\n")
        f.write(f"| :---: | :--- | :---: |\n")
        for s in steps:
            status_emoji = "✅ SUCCESS" if s["status"] == "SUCCESS" else "❌ FAILED"
            f.write(f"| {s['index']} | `{s['action']}` | {status_emoji} |\n")
        f.write("\n")
        
        f.write(f"## 3. Log & Failure Analysis\n")
        f.write(f"| Diagnostic Metric | Analysis Results |\n")
        f.write(f"| :--- | :--- |\n")
        f.write(f"| **Analysis Summary** | {analysis_data['summary']} |\n")
        f.write(f"| **Failure Category** | {analysis_data['category']} |\n")
        f.write(f"| **Diagnostic Details** | {analysis_data['details']} |\n")
        links_str = ", ".join([f"[{l}]({l})" for l in analysis_data["links"]]) if analysis_data["links"] else "None"
        f.write(f"| **Maestro Cloud Links** | {links_str} |\n\n")
        
        f.write(f"## 4. Raw Console Logs\n")
        f.write(f"```text\n")
        for line in log_lines:
            f.write(f"{line}\n")
        f.write(f"```\n")

def set_cell_background(cell, fill_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def write_docx_report(filepath, flow_name, device_id, verdict, steps, error_details, duration_sec, timestamp_str, log_lines):
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Base colors
    c_primary = RGBColor(15, 118, 110) # Teal
    c_dark = RGBColor(31, 41, 55)
    c_light = RGBColor(107, 114, 128)
    c_green = RGBColor(16, 185, 129)
    c_red = RGBColor(239, 68, 68)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("Konn3ct Mobile UI Test Report")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = c_primary
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run(f"Automated UI Functional Assessment  |  {timestamp_str}")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = c_light
    
    # Heading 1: Executive Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run_h1 = h1.add_run("1. Executive Summary")
    run_h1.font.name = "Arial"
    run_h1.font.size = Pt(16)
    run_h1.font.bold = True
    run_h1.font.color.rgb = c_primary
    
    # Summary Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    summary_data = [
        ("Test Flow File", flow_name),
        ("Target Device / Emulator", device_id),
        ("Execution Timestamp", timestamp_str),
        ("Total Duration", f"{duration_sec:.2f} seconds"),
        ("Overall Test Verdict", verdict)
    ]
    
    for idx, (label, val) in enumerate(summary_data):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        # Label cell
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10)
        
        # Value cell
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(10)
        
        if label == "Overall Test Verdict":
            r_val.font.bold = True
            r_val.font.color.rgb = c_green if verdict == "PASS" else r_val.font.color.rgb
            if verdict == "FAIL":
                r_val.font.color.rgb = c_red
                
        set_cell_background(cell_lbl, "F3F4F6")
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Heading 2: Step Execution Breakdown
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(8)
    run_h2 = h2.add_run("2. Step Execution Breakdown")
    run_h2.font.name = "Arial"
    run_h2.font.size = Pt(16)
    run_h2.font.bold = True
    run_h2.font.color.rgb = c_primary
    
    # Steps Table
    step_table = doc.add_table(rows=len(steps) + 1, cols=3)
    step_table.style = 'Light Shading Accent 1'
    
    # Headers
    headers = ["Step", "Action / Interaction Command", "Status"]
    for col_idx, text in enumerate(headers):
        cell = step_table.rows[0].cells[col_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F766E")
        set_cell_margins(cell)
        
    for r_idx, s in enumerate(steps):
        row = step_table.rows[r_idx + 1]
        c_idx, c_act, c_stat = row.cells[0], row.cells[1], row.cells[2]
        
        p0 = c_idx.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(str(s["index"]))
        r0.font.size = Pt(9)
        
        p1 = c_act.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(s["action"])
        r1.font.name = "Courier New"
        r1.font.size = Pt(9.5)
        
        p2 = c_stat.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(s["status"])
        r2.font.bold = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = c_green if s["status"] == "SUCCESS" else c_red
        
        # Zebra striping
        fill_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for cell in [c_idx, c_act, c_stat]:
            set_cell_background(cell, fill_color)
            set_cell_margins(cell)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Heading 3: Log & Failure Analysis
    h_analysis = doc.add_paragraph()
    h_analysis.paragraph_format.space_before = Pt(18)
    h_analysis.paragraph_format.space_after = Pt(8)
    run_h_analysis = h_analysis.add_run("3. Log & Failure Analysis")
    run_h_analysis.font.name = "Arial"
    run_h_analysis.font.size = Pt(16)
    run_h_analysis.font.bold = True
    run_h_analysis.font.color.rgb = c_primary
    
    analysis_data = analyze_logs(log_lines)
    
    ana_table = doc.add_table(rows=4, cols=2)
    ana_table.style = 'Light Shading Accent 1'
    
    ana_rows = [
        ("Analysis Summary", analysis_data["summary"]),
        ("Failure Category", analysis_data["category"]),
        ("Diagnostic Details", analysis_data["details"]),
        ("Maestro Cloud Links", "\n".join(analysis_data["links"]) if analysis_data["links"] else "None")
    ]
    
    for idx, (label, val) in enumerate(ana_rows):
        row = ana_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10)
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(10)
        
        if label == "Failure Category" and analysis_data["category"] != "Successful Run":
            r_val.font.bold = True
            r_val.font.color.rgb = c_red
        elif label == "Failure Category" and analysis_data["category"] == "Successful Run":
            r_val.font.bold = True
            r_val.font.color.rgb = c_green
            
        set_cell_background(cell_lbl, "F3F4F6")
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Heading 4: Raw Console Logs
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(8)
    run_h4 = h4.add_run("4. Raw Console Logs")
    run_h4.font.name = "Arial"
    run_h4.font.size = Pt(16)
    run_h4.font.bold = True
    run_h4.font.color.rgb = c_primary
    
    # Add raw logs as preformatted block
    p_log = doc.add_paragraph()
    p_log.paragraph_format.left_indent = Inches(0.2)
    p_log.paragraph_format.right_indent = Inches(0.2)
    p_log.paragraph_format.space_after = Pt(2)
    
    raw_text = "\n".join(log_lines)
    run_log = p_log.add_run(raw_text)
    run_log.font.name = "Courier New"
    run_log.font.size = Pt(8.5)
    run_log.font.color.rgb = RGBColor(75, 85, 99)
    
    doc.save(filepath)
